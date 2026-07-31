"""Convert GuildMaster_Player_Guide.md to DOCX with player-friendly formatting."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import re

doc = Document()

# ---- Page setup ----
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ---- Styles ----
style = doc.styles['Normal']
font = style.font
font.name = 'Segoe UI'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.2

for i in range(1, 4):
    h_style = doc.styles[f'Heading {i}']
    h_style.font.color.rgb = RGBColor(0xC0, 0x50, 0x20) if i == 1 else RGBColor(0x1A, 0x3C, 0x6E)
    h_style.font.bold = True

doc.styles['Heading 1'].font.size = Pt(22)
doc.styles['Heading 2'].font.size = Pt(16)
doc.styles['Heading 3'].font.size = Pt(13)

# Table style for tables
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

# Read markdown
with open('D:/Tinh/Rebuild_GuildMaster/Docs/GuildMaster_Player_Guide.md', 'r', encoding='utf-8') as f:
    content = f.read()

lines = content.split('\n')
in_code_block = False
table_data = []
in_table = False

for line in lines:
    # Code block
    if line.startswith('```'):
        in_code_block = not in_code_block
        continue
    if in_code_block:
        p = doc.add_paragraph(line)
        p.style = doc.styles['Normal']
        p.paragraph_format.left_indent = Cm(1)
        for run in p.runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        continue
    
    # Skip TOC entries
    if line.strip().startswith('- [') and '](' in line:
        continue
    if line.strip() == '---':
        continue
    
    # Headings
    if line.startswith('# ') and not line.startswith('## '):
        # Skip main title (handled as cover)
        continue
    if line.startswith('## '):
        doc.add_heading(line[3:], level=2)
        continue
    if line.startswith('### '):
        doc.add_heading(line[3:], level=3)
        continue
    
    # Table handling
    if line.strip().startswith('|') and '|' in line:
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if re.match(r'^\|[- :]+\|$', line):
            continue
        table_data.append(cells)
        in_table = True
        continue
    elif in_table and line.strip() == '':
        # Render table
        if table_data:
            rows = len(table_data)
            cols = max(len(r) for r in table_data) if table_data else 1
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Light Grid Accent 1'
            for i, row_data in enumerate(table_data):
                for j, cell_text in enumerate(row_data):
                    if j < len(table.rows[i].cells):
                        cell = table.rows[i].cells[j]
                        cell.text = cell_text
                        if i == 0:
                            set_cell_shading(cell, '1A3C6E')
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                    run.bold = True
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            if i % 2 == 0:
                                set_cell_shading(cell, 'F0F4FA')
            doc.add_paragraph()
        table_data = []
        in_table = False
        continue
    elif in_table and line.strip():
        continue
    
    table_data = []
    in_table = False
    
    # Regular paragraphs
    stripped = line.strip()
    if stripped == '':
        doc.add_paragraph()
        continue
    
    # List items
    if stripped.startswith('- **['):
        # Emoji headers
        doc.add_paragraph(stripped, style='List Bullet')
    elif stripped.startswith('- '):
        p = doc.add_paragraph(stripped[2:], style='List Bullet')
    elif stripped.startswith('* '):
        p = doc.add_paragraph(stripped[2:], style='List Bullet')
    elif re.match(r'^\d+\. ', stripped):
        p = doc.add_paragraph(stripped, style='List Number')
    elif stripped.startswith('|'):
        pass  # skip orphaned table lines
    else:
        # Handle inline bold and emoji
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*)', stripped)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part:
                run = p.add_run(part)

# ---- Cover Page ----
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml import parse_xml
# We insert cover as a new section at the beginning
section = doc.sections[0]
# Add page number footer
footer = section.footer
footer.is_linked_to_previous = False
p = footer.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Trang ')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# Save  
output_path = 'D:/Tinh/Rebuild_GuildMaster/Docs/GuildMaster_Player_Guide.docx'
doc.save(output_path)
print(f'DOCX saved to: {output_path}')
