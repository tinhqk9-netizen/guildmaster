"""Convert GuildMaster_Logic_Architecture.md to DOCX with proper formatting."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

doc = Document()

# ---- Page setup ----
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ---- Styles ----
style = doc.styles['Normal']
font = style.font
font.name = 'Segoe UI'
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for i in range(1, 4):
    h_style = doc.styles[f'Heading {i}']
    h_style.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    h_style.font.bold = True

doc.styles['Heading 1'].font.size = Pt(18)
doc.styles['Heading 2'].font.size = Pt(14)
doc.styles['Heading 3'].font.size = Pt(12)

# Code style
code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = 'Consolas'
code_style.font.size = Pt(9)
code_style.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
code_style.paragraph_format.space_before = Pt(2)
code_style.paragraph_format.space_after = Pt(2)
code_style.paragraph_format.left_indent = Inches(0.3)

# Read markdown
with open('D:/Tinh/Rebuild_GuildMaster/Docs/GuildMaster_Logic_Architecture.md', 'r', encoding='utf-8') as f:
    content = f.read()

# Process markdown
lines = content.split('\n')
in_code_block = False
code_buffer = []
table_mode = False

for line in lines:
    # Code block
    if line.startswith('```'):
        if in_code_block:
            # End code block
            for c_line in code_buffer:
                doc.add_paragraph(c_line, style='CodeBlock')
            code_buffer = []
            in_code_block = False
        else:
            in_code_block = True
            code_buffer = []
        continue
    
    if in_code_block:
        code_buffer.append(line)
        continue
    
    # Skip TOC entries
    if line.strip().startswith('- [') and '](' in line and '#' in line:
        continue
    
    # Skip the separator line
    if line.strip() == '---':
        continue
    
    # Headings
    if line.startswith('## '):
        doc.add_heading(line[3:], level=2)
        continue
    if line.startswith('### '):
        doc.add_heading(line[3:], level=3)
        continue
    if line.startswith('# ') and not line.startswith('# '):
        continue
    
    # Process horizontal rules and table separators
    if re.match(r'^\|[- |]+\|$', line):
        continue
    
    # Table rows
    if line.startswith('|') and '|' in line:
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if not table_mode:
            table_mode = True
            # Create table
            table = doc.add_table(rows=1, cols=len(cells))
            table.style = 'Light Grid Accent 1'
            for i, cell_text in enumerate(cells):
                table.rows[0].cells[i].text = cell_text
                for paragraph in table.rows[0].cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            prev_cells = cells
        else:
            row = table.add_row()
            for i, cell_text in enumerate(cells):
                if i < len(row.cells):
                    row.cells[i].text = cell_text
        continue
    
    if line.strip() == '' and table_mode:
        table_mode = False
        doc.add_paragraph()  # Add spacing after table
        continue
    
    # Regular paragraphs - process inline formatting
    if line.strip():
        # Check for list items
        stripped = line.strip()
        if stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(stripped[2:], style='List Bullet')
        elif re.match(r'^\d+\. ', stripped):
            p = doc.add_paragraph(stripped, style='List Number')
        else:
            # Bold patterns: **text**
            # Inline code: `text`
            parts = re.split(r'(\*\*.*?\*\*|`.*?`)', stripped)
            p = doc.add_paragraph()
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                else:
                    p.add_run(part)
    else:
        if not table_mode:
            doc.add_paragraph()

# ---- Cover page ----
# Insert cover at beginning
cover = Document()
for section in cover.sections:
    section.top_margin = Inches(2)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title
p = cover.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Guild Master')
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

p = cover.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Hệ Thống Logic Game — Toàn Tập')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x4A, 0x6F, 0xA5)

p = cover.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\nTài liệu diễn giải kiến trúc, luồng dữ liệu, công thức\nvà toàn bộ logic của dự án Rebuild Unity')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.italic = True

p = cover.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\nDựa trên code C# + DAD decompile từ APK gốc\nPhiên bản tài liệu 1.0')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

cover.add_page_break()

# Merge cover with content
for element in cover.element.body:
    doc.element.body.insert(0, element)

# ---- Add page numbers ----
from docx.oxml.ns import qn
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Trang ')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    # PAGE field
    fldChar1 = run._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run._element.addnext(fldChar1)
    instrText = run._element.makeelement(qn('w:instrText'), {})
    instrText.text = ' PAGE '
    fldChar1.addnext(instrText)
    fldChar2 = run._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    instrText.addnext(fldChar2)

# Save
output_path = 'D:/Tinh/Rebuild_GuildMaster/Docs/GuildMaster_Logic_Architecture.docx'
doc.save(output_path)
print(f'DOCX saved to: {output_path}')
